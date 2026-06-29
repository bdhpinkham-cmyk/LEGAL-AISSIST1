"""
ingestion.py
============
File ingestion + document production helpers used by the agentic engines.

Responsibilities
----------------
* Text extraction from PDFs and plain-text files.
* EXIF / metadata extraction from images (the "Metadata Extractor" toolkit item)
  to help authenticate the timeline.
* Audio / video transcription via the Deepgram or OpenAI Whisper REST APIs.
* PII auto-redaction (SSNs, phone numbers, emails, and configurable minor names).
* Court-ready DOCX export onto numbered pleading paper (lines 1-28 + caption).

Every function degrades gracefully: if an optional dependency is missing it
raises a clear, user-facing message rather than crashing the UI thread.
"""

from __future__ import annotations

import io
import os
import re
import time
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import requests

import config


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text(path: str) -> str:
    """Extract text from a supported file (PDF / TXT / MD). Returns ''."""
    lower = path.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(path)
    if lower.endswith((".txt", ".md", ".csv", ".log")):
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                return fh.read()
        except OSError as exc:
            raise IngestionError(f"Could not read {os.path.basename(path)}: {exc}")
    # Unknown text type: attempt a best-effort decode.
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()
    except OSError:
        return ""


def _extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise IngestionError("The 'pypdf' package is required to read PDFs.") from exc
    try:
        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()
    except Exception as exc:  # noqa: BLE001
        raise IngestionError(f"Failed to parse PDF: {exc}") from exc


class IngestionError(Exception):
    """User-facing ingestion failure."""


# ---------------------------------------------------------------------------
# Multimodal extraction (Gemini) — OCR + image description for big PDFs/images
# ---------------------------------------------------------------------------
_PDF_EXTRACT_PROMPT = (
    "You are a document-extraction engine for legal evidence. Extract ALL content "
    "from this PDF, page by page, in reading order:\n"
    "1. Transcribe every word of text verbatim — including text inside scanned "
    "pages, screenshots, stamps, handwriting, and form fields (perform OCR).\n"
    "2. For every photograph, diagram, chart, signature, or exhibit image, add a "
    "line beginning 'IMAGE:' that objectively describes what it depicts (people, "
    "objects, text visible, timestamps).\n"
    "Prefix each page with 'PAGE n'. Do not summarize, omit, or editorialize."
)
_IMAGE_EXTRACT_PROMPT = (
    "You are extracting a single image of legal evidence. Output two sections:\n"
    "TEXT: every word visible in the image, transcribed verbatim (OCR). If none, "
    "write 'TEXT: (none)'.\n"
    "DESCRIPTION: an objective description of what the image depicts — people, "
    "objects, setting, any visible timestamps or identifying detail. Do not "
    "speculate or draw legal conclusions."
)

_PDF_MIME = "application/pdf"
_IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
}


def extract_document(
    path: str,
    filename: str,
    extraction: tuple,
    log=lambda _m: None,
) -> Tuple[str, List[Dict]]:
    """Extract text from a file, using multimodal OCR for PDFs/images.

    ``extraction`` is the (provider, model, api_key) tuple for the extraction
    tier (typically Gemini). PDFs and images are sent to the multimodal model so
    that scanned pages and embedded photos are OCR'd and described. If the
    multimodal path is unavailable (no Gemini key, or a non-Gemini extraction
    provider), PDFs fall back to local ``pypdf`` text extraction and images fall
    back to an EXIF description.

    Returns ``(text, gaps)`` — ``gaps`` describes any page range that could not
    be extracted even after retries, so callers can surface a visible coverage
    report instead of silently losing content.
    """
    import gemini_client

    provider, model, key = (extraction + ("", "", ""))[:3]
    gemini_ready = provider == "gemini" and gemini_client.is_configured()
    lower = filename.lower()
    ext = os.path.splitext(lower)[1]

    if ext == ".pdf":
        if gemini_ready:
            try:
                text, gaps = gemini_extract_pdf(path, model, key, log=log)
                if text.strip():
                    return text, gaps
                log("Multimodal extraction returned nothing; falling back to text layer.")
            except IngestionError as exc:
                log(f"Multimodal PDF extraction failed ({exc}); using local text layer.")
        return _extract_pdf(path), []

    if ext in _IMAGE_MIME:
        if gemini_ready:
            try:
                return gemini_extract_image(path, model, key), []
            except IngestionError as exc:
                log(f"Multimodal image extraction failed ({exc}).")
        meta = extract_image_metadata(path)
        return "Image evidence (no multimodal model configured).\n" + "\n".join(
            f"{k}: {v}" for k, v in meta.items()
        ), []

    return extract_text(path), []


def gemini_extract_image(path: str, model: str, key: str = "") -> str:
    """OCR + describe a single image via Gemini (AI Studio or Vertex AI)."""
    import gemini_client

    ext = os.path.splitext(path.lower())[1]
    mime = _IMAGE_MIME.get(ext, "image/jpeg")
    try:
        return gemini_client.generate_from_file(
            model, path, mime, _IMAGE_EXTRACT_PROMPT,
            max_tokens=config.GEMINI_EXTRACT_MAX_TOKENS,
        )
    except gemini_client.GeminiError as exc:
        raise IngestionError(str(exc)) from exc


def gemini_extract_pdf(
    path: str, model: str, key: str = "", log=lambda _m: None
) -> Tuple[str, List[Dict]]:
    """Extract a (potentially huge, image-heavy) PDF via Gemini, page-batched.

    The PDF is split into ``PDF_PAGE_BATCH``-page chunks so each multimodal call
    stays under the model's output-token cap even for dense scanned pages, and so
    each inline payload stays small (works on both AI Studio and Vertex). Every
    batch's text is concatenated, then the caller chunks it into the RAG store.

    Each batch is retried (``PDF_BATCH_MAX_RETRIES``, exponential backoff)
    before being given up on. A batch that still fails leaves a visible
    ``FAILED TO EXTRACT`` marker in the returned text and a matching entry in
    the returned ``gaps`` list — nothing is dropped silently. Pages beyond the
    ``PDF_MAX_BATCHES`` cap are likewise recorded as a gap rather than an
    unflagged truncation.
    """
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:  # pragma: no cover
        raise IngestionError("The 'pypdf' package is required for PDF batching.") from exc

    import gemini_client

    try:
        reader = PdfReader(path)
        n_pages = len(reader.pages)
    except Exception as exc:  # noqa: BLE001
        raise IngestionError(f"Could not open PDF: {exc}") from exc

    batch = max(1, config.PDF_PAGE_BATCH)
    n_batches = (n_pages + batch - 1) // batch
    gaps: List[Dict] = []
    if n_batches > config.PDF_MAX_BATCHES:
        skipped_start = config.PDF_MAX_BATCHES * batch + 1
        log(
            f"PDF has {n_pages} pages; processing the first "
            f"{config.PDF_MAX_BATCHES * batch} pages."
        )
        gaps.append({
            "start_page": skipped_start,
            "end_page": n_pages,
            "reason": f"exceeded PDF_MAX_BATCHES cap ({config.PDF_MAX_BATCHES} batches)",
        })
        n_batches = config.PDF_MAX_BATCHES

    log(f"Parsing {n_pages}-page PDF in {n_batches} batch(es) of {batch} pages…")
    out_parts: List[str] = []
    succeeded = 0

    for b in range(n_batches):
        start = b * batch
        end = min(start + batch, n_pages)
        log(f"  Extracting pages {start + 1}-{end} (batch {b + 1}/{n_batches})…")
        tmp_path = _write_pdf_subset(reader, PdfWriter, start, end)
        prompt = (
            f"{_PDF_EXTRACT_PROMPT}\n\n(This is pages {start + 1}-{end} of a "
            f"{n_pages}-page document.)"
        )
        text = ""
        last_exc: Optional[Exception] = None
        try:
            for attempt in range(config.PDF_BATCH_MAX_RETRIES):
                try:
                    text = gemini_client.generate_from_file(
                        model, tmp_path, _PDF_MIME, prompt,
                        max_tokens=config.GEMINI_EXTRACT_MAX_TOKENS,
                    )
                    last_exc = None
                    break
                except gemini_client.GeminiError as exc:
                    last_exc = exc
                    if attempt < config.PDF_BATCH_MAX_RETRIES - 1:
                        log(f"  Batch {b + 1} attempt {attempt + 1} failed: {exc} (retrying)")
                        time.sleep(config.RETRY_BASE_DELAY * (2 ** attempt))
        finally:
            _safe_unlink(tmp_path)

        if last_exc is not None:
            reason = str(last_exc)
            log(f"  Batch {b + 1} failed after {config.PDF_BATCH_MAX_RETRIES} attempt(s): {reason}")
            out_parts.append(f"=== PAGES {start + 1}-{end} FAILED TO EXTRACT ({reason}) ===")
            gaps.append({"start_page": start + 1, "end_page": end, "reason": reason})
        else:
            succeeded += 1
            if text:
                out_parts.append(f"=== PAGES {start + 1}-{end} ===\n{text}")

    log(f"PDF parsed: {succeeded}/{n_batches} batches succeeded.")
    return "\n\n".join(out_parts).strip(), gaps


def _write_pdf_subset(reader, writer_cls, start: int, end: int) -> str:
    import tempfile

    config.ensure_directories()
    writer = writer_cls()
    for i in range(start, end):
        writer.add_page(reader.pages[i])
    fd, tmp_path = tempfile.mkstemp(suffix=".pdf", dir=str(config.DATA_ROOT))
    os.close(fd)
    with open(tmp_path, "wb") as fh:
        writer.write(fh)
    return tmp_path


def _safe_unlink(path: str) -> None:
    try:
        os.unlink(path)
    except OSError:
        pass


# ---------------------------------------------------------------------------
# Image metadata (EXIF) extraction
# ---------------------------------------------------------------------------
def extract_image_metadata(path: str) -> Dict[str, str]:
    """Return a flat dict of useful EXIF fields (timestamp, GPS, camera)."""
    try:
        from PIL import Image, ExifTags
    except ImportError as exc:  # pragma: no cover
        raise IngestionError("The 'Pillow' package is required to read image metadata.") from exc

    meta: Dict[str, str] = {}
    try:
        with Image.open(path) as img:
            meta["format"] = img.format or ""
            meta["size"] = f"{img.width}x{img.height}"
            exif = img.getexif()
            if not exif:
                return meta
            tag_map = {v: k for k, v in ExifTags.TAGS.items()}
            for tag_id, value in exif.items():
                name = ExifTags.TAGS.get(tag_id, str(tag_id))
                if name in ("DateTime", "DateTimeOriginal", "Make", "Model", "Software"):
                    meta[name] = str(value)
            # GPS sub-IFD
            gps_ifd_id = tag_map.get("GPSInfo")
            if gps_ifd_id and gps_ifd_id in exif:
                gps = exif.get_ifd(gps_ifd_id)
                coords = _parse_gps(gps)
                if coords:
                    meta["GPSLatitude"], meta["GPSLongitude"] = coords
    except Exception as exc:  # noqa: BLE001
        meta["error"] = str(exc)
    return meta


def _parse_gps(gps: Dict) -> Optional[Tuple[str, str]]:
    try:
        from PIL import ExifTags

        gps_tags = {ExifTags.GPSTAGS.get(k, k): v for k, v in gps.items()}

        def _to_deg(values, ref) -> float:
            d, m, s = [float(x) for x in values]
            deg = d + m / 60.0 + s / 3600.0
            if ref in ("S", "W"):
                deg = -deg
            return deg

        lat = gps_tags.get("GPSLatitude")
        lat_ref = gps_tags.get("GPSLatitudeRef")
        lon = gps_tags.get("GPSLongitude")
        lon_ref = gps_tags.get("GPSLongitudeRef")
        if lat and lon and lat_ref and lon_ref:
            return (
                f"{_to_deg(lat, lat_ref):.6f}",
                f"{_to_deg(lon, lon_ref):.6f}",
            )
    except Exception:  # noqa: BLE001
        return None
    return None


# ---------------------------------------------------------------------------
# Audio / video transcription
# ---------------------------------------------------------------------------
def _chunk_boundaries(total_ms: int, chunk_ms: int) -> List[Tuple[int, int]]:
    """Pure helper: split [0, total_ms) into consecutive (start_ms, end_ms) spans."""
    if total_ms <= 0:
        return [(0, 0)]
    chunk_ms = max(1, chunk_ms)
    bounds: List[Tuple[int, int]] = []
    start = 0
    while start < total_ms:
        end = min(start + chunk_ms, total_ms)
        bounds.append((start, end))
        start = end
    return bounds


def _fmt_hms(seconds: float) -> str:
    seconds = int(max(0, seconds))
    h, rem = divmod(seconds, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def split_audio(
    path: str,
    chunk_seconds: int = config.AUDIO_CHUNK_SECONDS,
    log=lambda _m: None,
) -> List[Tuple[str, float, float]]:
    """Slice a large audio/video file into ``chunk_seconds`` segments.

    Returns a list of ``(chunk_path, start_seconds, end_seconds)``. Files
    already under ``AUDIO_SIZE_LIMIT_BYTES`` are returned unchanged as a single
    segment — today's whole-file behavior is preserved exactly, and pydub is
    never imported for the common case of short recordings. Raises
    ``ImportError`` (propagated, not swallowed) if pydub isn't installed and
    the file is actually large enough to require chunking; callers fall back
    to whole-file transcription on that error.
    """
    try:
        size = os.path.getsize(path)
    except OSError:
        size = 0

    if size and size <= config.AUDIO_SIZE_LIMIT_BYTES:
        return [(path, 0.0, 0.0)]

    from pydub import AudioSegment
    import tempfile

    audio = AudioSegment.from_file(path)
    total_ms = len(audio)
    chunk_ms = max(1000, int(chunk_seconds * 1000))
    bounds = _chunk_boundaries(total_ms, chunk_ms)

    if len(bounds) <= 1:
        return [(path, 0.0, total_ms / 1000.0)]

    if len(bounds) > config.AUDIO_MAX_CHUNKS:
        log(
            f"Audio is {_fmt_hms(total_ms / 1000.0)} long; only processing the first "
            f"{config.AUDIO_MAX_CHUNKS} chunk(s) (~{_fmt_hms(config.AUDIO_MAX_CHUNKS * chunk_seconds)})."
        )
        bounds = bounds[: config.AUDIO_MAX_CHUNKS]

    config.ensure_directories()
    chunks: List[Tuple[str, float, float]] = []
    for start_ms, end_ms in bounds:
        segment = audio[start_ms:end_ms]
        fd, tmp_path = tempfile.mkstemp(suffix=".mp3", dir=str(config.DATA_ROOT))
        os.close(fd)
        segment.export(tmp_path, format="mp3")
        chunks.append((tmp_path, start_ms / 1000.0, end_ms / 1000.0))
    return chunks


def transcribe_audio(
    path: str,
    deepgram_key: str = "",
    openai_key: str = "",
    log=lambda _m: None,
) -> Tuple[str, List[Dict]]:
    """Transcribe an audio/video file. Prefers Deepgram, falls back to Whisper.

    Large files are sliced into ``AUDIO_CHUNK_SECONDS`` segments first (see
    ``split_audio``) so no single request exceeds the transcription API's size
    limits and a failed segment doesn't abort the whole recording. Each
    chunk's text is prefixed with a ``=== TIME hh:mm:ss-hh:mm:ss ===`` header.
    Returns ``(transcript, gaps)``; a chunk whose retries are exhausted is
    recorded as a gap rather than aborting the file.
    """
    if not deepgram_key and not openai_key:
        raise IngestionError(
            "Transcription requires a Deepgram or OpenAI API key (set one in Settings)."
        )

    def _transcribe_one(chunk_path: str) -> str:
        if deepgram_key:
            return _transcribe_deepgram(chunk_path, deepgram_key)
        return _transcribe_whisper(chunk_path, openai_key)

    try:
        chunks = split_audio(path, log=log)
    except ImportError:
        log("Audio chunking unavailable (pydub/ffmpeg not found); sending whole file.")
        chunks = [(path, 0.0, 0.0)]

    if len(chunks) == 1 and chunks[0][0] == path:
        return _transcribe_one(path), []

    log(f"Transcribing audio in {len(chunks)} chunk(s)…")
    out_parts: List[str] = []
    gaps: List[Dict] = []
    for i, (chunk_path, start_s, end_s) in enumerate(chunks):
        log(f"  Transcribing {_fmt_hms(start_s)}-{_fmt_hms(end_s)} (chunk {i + 1}/{len(chunks)})…")
        try:
            text = _transcribe_one(chunk_path)
            if text:
                out_parts.append(f"=== TIME {_fmt_hms(start_s)}-{_fmt_hms(end_s)} ===\n{text}")
        except IngestionError as exc:
            log(f"  Chunk {i + 1} failed: {exc}")
            gaps.append({"start_s": start_s, "end_s": end_s, "reason": str(exc)})
        finally:
            if chunk_path != path:
                _safe_unlink(chunk_path)

    log(f"Audio parsed: {len(chunks) - len(gaps)}/{len(chunks)} chunks succeeded.")
    return "\n\n".join(out_parts).strip(), gaps


def _transcribe_deepgram(path: str, key: str) -> str:
    url = "https://api.deepgram.com/v1/listen?smart_format=true&punctuate=true"
    headers = {"Authorization": f"Token {key}", "Content-Type": "application/octet-stream"}
    last_exc: Optional[Exception] = None
    for attempt in range(config.MAX_API_RETRIES):
        try:
            with open(path, "rb") as fh:
                resp = requests.post(
                    url, headers=headers, data=fh, timeout=300
                )
            if resp.status_code == 429:
                raise RuntimeError("rate limit")
            resp.raise_for_status()
            data = resp.json()
            return (
                data["results"]["channels"][0]["alternatives"][0]["transcript"].strip()
            )
        except Exception as exc:  # noqa: BLE001
            last_exc = exc
            if attempt < config.MAX_API_RETRIES - 1:
                time.sleep(config.RETRY_BASE_DELAY * (2 ** attempt))
    raise IngestionError(f"Deepgram transcription failed: {last_exc}")


def _transcribe_whisper(path: str, key: str) -> str:
    url = "https://api.openai.com/v1/audio/transcriptions"
    headers = {"Authorization": f"Bearer {key}"}
    last_exc: Optional[Exception] = None
    for attempt in range(config.MAX_API_RETRIES):
        try:
            with open(path, "rb") as fh:
                files = {"file": (os.path.basename(path), fh)}
                data = {"model": "whisper-1"}
                resp = requests.post(
                    url, headers=headers, files=files, data=data, timeout=300
                )
            if resp.status_code == 429:
                raise RuntimeError("rate limit")
            resp.raise_for_status()
            return resp.json().get("text", "").strip()
        except Exception as exc:  # noqa: BLE001
            last_exc = exc
            if attempt < config.MAX_API_RETRIES - 1:
                time.sleep(config.RETRY_BASE_DELAY * (2 ** attempt))
    raise IngestionError(f"Whisper transcription failed: {last_exc}")


# ---------------------------------------------------------------------------
# PII redaction
# ---------------------------------------------------------------------------
_SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
_PHONE_RE = re.compile(r"\b(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b")
_EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
_CC_RE = re.compile(r"\b(?:\d[ -]*?){13,16}\b")


def redact_pii(text: str, minor_names: Optional[List[str]] = None) -> Tuple[str, List[str]]:
    """Scrub PII from ``text``. Returns (redacted_text, list_of_redaction_notes)."""
    notes: List[str] = []
    redacted = text

    def _sub(pattern: re.Pattern, label: str, value: str) -> None:
        nonlocal redacted
        count = len(pattern.findall(redacted))
        if count:
            redacted = pattern.sub(value, redacted)
            notes.append(f"Redacted {count} {label}")

    _sub(_SSN_RE, "SSN(s)", "[REDACTED-SSN]")
    _sub(_CC_RE, "possible card number(s)", "[REDACTED-NUMBER]")
    _sub(_PHONE_RE, "phone number(s)", "[REDACTED-PHONE]")
    _sub(_EMAIL_RE, "email address(es)", "[REDACTED-EMAIL]")

    for name in minor_names or []:
        name = name.strip()
        if not name:
            continue
        pattern = re.compile(re.escape(name), re.IGNORECASE)
        count = len(pattern.findall(redacted))
        if count:
            redacted = pattern.sub("[REDACTED-MINOR]", redacted)
            notes.append(f"Redacted {count} reference(s) to a protected name")

    return redacted, notes


# ---------------------------------------------------------------------------
# Court-ready DOCX export (numbered pleading paper)
# ---------------------------------------------------------------------------
def export_pleading_docx(
    out_path: str,
    *,
    party_name: str,
    court: str,
    case_number: str,
    judge: str,
    title: str,
    body: str,
    plaintiff: str = "THE PEOPLE / PLAINTIFF",
    defendant: str = "",
) -> str:
    """Render a legal pleading to DOCX with a caption and 1-28 line numbers.

    Line numbering on pleading paper is produced by enabling Word's built-in
    line-number field in the section properties (restarts each page), which is
    how real pleading templates do it — this keeps numbering correct regardless
    of how the text reflows. A numbered left column is also drawn for on-screen
    fidelity.
    """
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.shared import Pt, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:  # pragma: no cover
        raise IngestionError("The 'python-docx' package is required for DOCX export.") from exc

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Enable Word line numbering (continuous, restart each page) on the section.
    sectPr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:restart"), "newPage")
    ln.set(qn("w:distance"), "360")
    sectPr.append(ln)

    base = doc.styles["Normal"]
    base.font.name = "Times New Roman"
    base.font.size = Pt(12)

    # Party / attorney block.
    head = doc.add_paragraph()
    head.add_run(f"{party_name}\nDefendant In Pro Per\n").bold = False

    # Court caption.
    court_p = doc.add_paragraph()
    run = court_p.add_run(court.upper() if court else "SUPERIOR COURT OF THE STATE")
    run.bold = True
    court_p.alignment = 1  # centre

    # Caption table: parties vs case number.
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.text = (
        f"{plaintiff},\n        Plaintiff,\n   vs.\n"
        f"{defendant or party_name},\n        Defendant."
    )
    right.text = (
        f"Case No.: {case_number or '__________'}\n\n"
        f"{title.upper()}\n\n"
        f"Judge: {judge or '__________'}"
    )

    doc.add_paragraph("")

    # Body with an explicit numbered column for on-screen pleading fidelity.
    line_no = 1
    for raw_line in body.split("\n"):
        # Wrap long lines so the numbered column stays readable.
        for segment in _wrap(raw_line, 95) or [""]:
            p = doc.add_paragraph()
            num_run = p.add_run(f"{line_no:>2}  ")
            num_run.font.size = Pt(9)
            p.add_run(segment)
            line_no += 1
            if line_no > 28:
                line_no = 1

    config.ensure_directories()
    doc.save(out_path)
    return out_path


def _wrap(text: str, width: int) -> List[str]:
    if not text:
        return [""]
    words = text.split(" ")
    lines: List[str] = []
    current = ""
    for w in words:
        if len(current) + len(w) + 1 > width and current:
            lines.append(current)
            current = w
        else:
            current = f"{current} {w}".strip()
    if current:
        lines.append(current)
    return lines
